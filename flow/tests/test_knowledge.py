# Copyright (c) 2026, Frappe Technologies and contributors
# License: MIT. See LICENSE

import json
from types import SimpleNamespace
from unittest.mock import patch

import frappe
from frappe.tests import IntegrationTestCase

from flow.knowledge import Knowledge, store
from flow.knowledge.chunker import chunk_text
from flow.knowledge.embedder import embed_texts, probe_dimension
from flow.knowledge.extract import (
	_extract_docx,
	_extract_html,
	_extract_pdf,
	_extract_xlsx,
	_validate_public_url,
	extract,
)
from flow.knowledge.ingest import (
	ingest_source,
	purge_source,
	reconcile_source,
	sync_due_sources,
)
from flow.knowledge.retriever import retrieve

DIM = 4


def _rows():
	return [
		{
			"id": 1,
			"kb": "Helpdesk",
			"source": "SRC-a",
			"content": "wifi not connecting on macbook",
			"vector": [1.0, 0.0, 0.0, 0.0],
		},
		{
			"id": 2,
			"kb": "Helpdesk",
			"source": "SRC-a",
			"content": "printer not detected on windows",
			"vector": [0.0, 1.0, 0.0, 0.0],
		},
		{
			"id": 3,
			"kb": "Onboarding",
			"source": "SRC-b",
			"content": "how to set up your laptop",
			"vector": [0.9, 0.1, 0.0, 0.0],
		},
	]


class TestKnowledgeStore(IntegrationTestCase):
	def setUp(self):
		store.drop_table()

	def tearDown(self):
		store.drop_table()

	def _seed(self):
		store.ensure_table_for_dimension(DIM)
		store.add(_rows())

	def test_ensure_table_is_idempotent(self):
		store.ensure_table_for_dimension(DIM)
		store.ensure_table_for_dimension(DIM)
		self.assertTrue(store.table_exists())
		self.assertEqual(store.table_dimension(), DIM)

	def test_ensure_table_rejects_dimension_change(self):
		store.ensure_table_for_dimension(DIM)
		with self.assertRaisesRegex(frappe.ValidationError, "Rebuild"):
			store.ensure_table_for_dimension(DIM + 1)

	def test_ensure_table_rejects_invalid_dimension(self):
		with self.assertRaises(ValueError):
			store.ensure_table_for_dimension(0)

	def test_search_before_table_exists_returns_empty(self):
		self.assertEqual(store.search([1.0, 0.0, 0.0, 0.0]), [])

	def test_search_on_empty_table_returns_empty(self):
		store.ensure_table_for_dimension(DIM)
		self.assertEqual(store.search([1.0, 0.0, 0.0, 0.0]), [])

	def test_vector_search_ranks_by_similarity(self):
		self._seed()
		hits = store.search([1.0, 0.0, 0.0, 0.0], limit=2)
		self.assertEqual([h["id"] for h in hits], [1, 3])
		self.assertGreater(hits[0]["score"], hits[1]["score"])

	def test_search_scopes_to_knowledge_bases(self):
		self._seed()
		hits = store.search([1.0, 0.0, 0.0, 0.0], kbs=["Onboarding"], limit=5)
		self.assertEqual([h["kb"] for h in hits], ["Onboarding"])

	def test_hybrid_search_surfaces_keyword_match(self):
		self._seed()
		hits = store.search([0.0, 0.0, 1.0, 0.0], text="printer", limit=3)
		self.assertIn(2, [h["id"] for h in hits])

	def test_hybrid_search_finds_rows_added_after_index_creation(self):
		self._seed()
		store.add(
			[
				{
					"id": 4,
					"kb": "Helpdesk",
					"source": "SRC-c",
					"content": "vpn conflict after update",
					"vector": [0.0, 0.0, 1.0, 0.0],
				}
			]
		)
		hits = store.search([0.0, 0.0, 1.0, 0.0], text="vpn", limit=3)
		self.assertEqual(hits[0]["id"], 4)

	def test_delete_by_source(self):
		self._seed()
		store.delete(source="SRC-a")
		hits = store.search([1.0, 0.0, 0.0, 0.0], limit=5)
		self.assertEqual([h["id"] for h in hits], [3])

	def test_delete_by_ids(self):
		self._seed()
		store.delete(ids=[1, 3])
		hits = store.search([1.0, 0.0, 0.0, 0.0], limit=5)
		self.assertEqual([h["id"] for h in hits], [2])

	def test_delete_with_empty_ids_is_noop(self):
		self._seed()
		store.delete(ids=[])
		self.assertEqual(len(store.search([1.0, 0.0, 0.0, 0.0], limit=5)), 3)

	def test_delete_requires_a_criterion(self):
		with self.assertRaises(ValueError):
			store.delete()

	def test_filter_values_with_quotes_are_escaped(self):
		self._seed()
		hits = store.search([1.0, 0.0, 0.0, 0.0], kbs=["x'; DROP TABLE chunks; --"], limit=5)
		self.assertEqual(hits, [])
		self.assertTrue(store.table_exists())

	def test_add_requires_table(self):
		with self.assertRaisesRegex(frappe.ValidationError, "not initialized"):
			store.add(_rows())


def _embedding_response(vectors, shuffle=False):
	data = [{"embedding": v, "index": i} for i, v in enumerate(vectors)]
	if shuffle:
		data = list(reversed(data))
	return SimpleNamespace(data=data)


def _make_model(title="Embed Model", model_id="openai/text-embedding-3-small"):
	if frappe.db.exists("AI Model", title):
		return frappe.get_doc("AI Model", title)
	return frappe.get_doc(
		{"doctype": "AI Model", "title": title, "model_id": model_id, "api_key": "sk-test", "enabled": 1}
	).insert()


def _set_settings(**values):
	for fieldname, value in values.items():
		frappe.db.set_single_value("AI Knowledge Settings", fieldname, value)
	frappe.clear_document_cache("AI Knowledge Settings", "AI Knowledge Settings")


class TestEmbedder(IntegrationTestCase):
	def setUp(self):
		self.model = _make_model()
		_set_settings(embedding_model=self.model.name)

	def tearDown(self):
		frappe.db.rollback()

	def test_embed_texts_returns_vectors_in_input_order(self):
		response = _embedding_response([[1.0, 2.0], [3.0, 4.0]], shuffle=True)
		with patch("litellm.embedding", return_value=response) as mocked:
			vectors = embed_texts(["a", "b"])
		self.assertEqual(vectors, [[1.0, 2.0], [3.0, 4.0]])
		kwargs = mocked.call_args.kwargs
		self.assertEqual(kwargs["model"], "openai/text-embedding-3-small")
		self.assertEqual(kwargs["api_key"], "sk-test")
		self.assertEqual(kwargs["input"], ["a", "b"])

	def test_embed_texts_batches_large_input(self):
		texts = [f"t{i}" for i in range(100)]

		def fake_embedding(input, **kwargs):
			return _embedding_response([[0.0]] * len(input))

		with patch("litellm.embedding", side_effect=fake_embedding) as mocked:
			vectors = embed_texts(texts)
		self.assertEqual(len(vectors), 100)
		self.assertEqual(mocked.call_count, 2)
		self.assertEqual(len(mocked.call_args_list[0].kwargs["input"]), 96)
		self.assertEqual(len(mocked.call_args_list[1].kwargs["input"]), 4)

	def test_embed_texts_empty_input_skips_api(self):
		with patch("litellm.embedding") as mocked:
			self.assertEqual(embed_texts([]), [])
		mocked.assert_not_called()

	def test_embed_texts_requires_configured_model(self):
		_set_settings(embedding_model="")
		with self.assertRaisesRegex(frappe.ValidationError, "Knowledge Settings"):
			embed_texts(["a"])

	def test_embed_texts_rejects_disabled_model(self):
		self.model.enabled = 0
		self.model.save()
		with self.assertRaisesRegex(frappe.ValidationError, "disabled"):
			embed_texts(["a"])

	def test_embed_texts_count_mismatch_throws(self):
		response = _embedding_response([[1.0]])
		with patch("litellm.embedding", return_value=response):
			with self.assertRaisesRegex(frappe.ValidationError, "2 inputs"):
				embed_texts(["a", "b"])

	def test_embed_provider_error_surfaces_as_validation_error(self):
		with patch("litellm.embedding", side_effect=RuntimeError("invalid api key")):
			with self.assertRaisesRegex(frappe.ValidationError, "invalid api key"):
				embed_texts(["a"])

	def test_probe_dimension(self):
		response = _embedding_response([[0.0] * 1536])
		with patch("litellm.embedding", return_value=response):
			self.assertEqual(probe_dimension(self.model.name), 1536)


class TestKnowledgeSettings(IntegrationTestCase):
	def setUp(self):
		self.model = _make_model()
		self.alt_model = _make_model(title="Embed Model Alt", model_id="openai/text-embedding-3-large")
		_set_settings(embedding_model="", embedding_dimension=0)

	def tearDown(self):
		frappe.db.rollback()

	def _save_settings(self, **values):
		settings = frappe.get_doc("AI Knowledge Settings")
		settings.update(values)
		settings.save()
		return settings

	def _make_chunk(self):
		kb = frappe.get_doc({"doctype": "AI Knowledge Base", "title": "Settings KB"}).insert()
		source = frappe.get_doc(
			{
				"doctype": "AI Knowledge Source",
				"knowledge_base": kb.name,
				"source_type": "Text",
				"title": "Settings Source",
				"content": "hello",
			}
		).insert()
		return frappe.get_doc(
			{
				"doctype": "AI Knowledge Chunk",
				"knowledge_base": kb.name,
				"source": source.name,
				"chunk_index": 0,
				"content": "hello",
			}
		).insert()

	def test_save_probes_dimension(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536) as probe:
			settings = self._save_settings(embedding_model=self.model.name)
		self.assertEqual(settings.embedding_dimension, 1536)
		probe.assert_called_once_with(self.model.name)

	def test_save_unchanged_model_does_not_reprobe(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536):
			self._save_settings(embedding_model=self.model.name)
		with patch("flow.knowledge.embedder.probe_dimension") as probe:
			self._save_settings(chunk_size=500)
		probe.assert_not_called()

	def test_model_change_reprobes_without_chunks(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536):
			self._save_settings(embedding_model=self.model.name)
		with patch("flow.knowledge.embedder.probe_dimension", return_value=3072):
			settings = self._save_settings(embedding_model=self.alt_model.name)
		self.assertEqual(settings.embedding_dimension, 3072)

	def test_model_change_blocked_while_chunks_exist(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536):
			self._save_settings(embedding_model=self.model.name)
		self._make_chunk()
		with self.assertRaisesRegex(frappe.ValidationError, "Delete all knowledge sources"):
			self._save_settings(embedding_model=self.alt_model.name)

	def test_clear_model_blocked_while_chunks_exist(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536):
			self._save_settings(embedding_model=self.model.name)
		self._make_chunk()
		with self.assertRaisesRegex(frappe.ValidationError, "Delete all knowledge sources"):
			self._save_settings(embedding_model="")

	def test_model_change_flag_bypasses_guard(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536):
			self._save_settings(embedding_model=self.model.name)
		self._make_chunk()
		settings = frappe.get_doc("AI Knowledge Settings")
		settings.embedding_model = self.alt_model.name
		settings.flags.allow_embedding_model_change = True
		with patch("flow.knowledge.embedder.probe_dimension", return_value=3072):
			settings.save()
		self.assertEqual(settings.embedding_dimension, 3072)

	def test_clear_model_resets_dimension(self):
		with patch("flow.knowledge.embedder.probe_dimension", return_value=1536):
			self._save_settings(embedding_model=self.model.name)
		settings = self._save_settings(embedding_model="")
		self.assertEqual(settings.embedding_dimension, 0)

	def test_chunk_overlap_must_be_smaller_than_chunk_size(self):
		with self.assertRaisesRegex(frappe.ValidationError, "Chunk Overlap"):
			self._save_settings(chunk_size=100, chunk_overlap=100)

	def test_chunk_size_must_be_positive(self):
		with self.assertRaisesRegex(frappe.ValidationError, "Chunk Size"):
			self._save_settings(chunk_size=0)


class TestChunker(IntegrationTestCase):
	def test_empty_text_returns_no_chunks(self):
		self.assertEqual(chunk_text("   \n  ", chunk_size=100, overlap=10), [])

	def test_short_text_is_a_single_chunk(self):
		self.assertEqual(chunk_text("hello world", chunk_size=100, overlap=10), ["hello world"])

	def test_long_text_splits_within_size_and_covers_all_words(self):
		words = [f"w{i}" for i in range(200)]
		text = " ".join(words)
		chunks = chunk_text(text, chunk_size=100, overlap=20)
		self.assertGreater(len(chunks), 1)
		self.assertTrue(all(len(c) <= 100 for c in chunks))
		joined = " ".join(chunks)
		self.assertTrue(all(w in joined for w in words))

	def test_consecutive_chunks_overlap(self):
		text = " ".join(f"token{i}" for i in range(100))
		chunks = chunk_text(text, chunk_size=80, overlap=30)
		tail = chunks[0].split()[-1]
		self.assertIn(tail, chunks[1])

	def test_text_without_whitespace_hard_splits(self):
		chunks = chunk_text("a" * 250, chunk_size=100, overlap=0)
		self.assertEqual([len(c) for c in chunks], [100, 100, 50])

	def test_invalid_chunk_size_raises(self):
		with self.assertRaises(ValueError):
			chunk_text("x", chunk_size=0, overlap=0)

	def test_overlap_not_smaller_than_size_raises(self):
		with self.assertRaises(ValueError):
			chunk_text("x", chunk_size=10, overlap=10)


class TestExtract(IntegrationTestCase):
	def tearDown(self):
		frappe.db.rollback()

	def test_text_source(self):
		source = frappe._dict(source_type="Text", content="  some knowledge  ")
		docs = extract(source)
		self.assertEqual(len(docs), 1)
		self.assertEqual(docs[0].text, "some knowledge")
		self.assertIsNone(docs[0].reference_name)

	def test_empty_text_source_yields_nothing(self):
		self.assertEqual(extract(frappe._dict(source_type="Text", content="   ")), [])

	def test_file_source_reads_text_file(self):
		file_doc = frappe.get_doc(
			{"doctype": "File", "file_name": "note.txt", "content": "file body", "is_private": 1}
		).insert()
		docs = extract(frappe._dict(source_type="File", file=file_doc.file_url))
		self.assertEqual(docs[0].text, "file body")

	def test_file_source_rejects_unsupported_format(self):
		file_doc = frappe.get_doc(
			{"doctype": "File", "file_name": "data.bin", "content": "x", "is_private": 1}
		).insert()
		with self.assertRaisesRegex(frappe.ValidationError, "Unsupported file format"):
			extract(frappe._dict(source_type="File", file=file_doc.file_url))

	def test_doctype_source_one_record_per_document(self):
		frappe.get_doc({"doctype": "ToDo", "description": "first task", "status": "Open"}).insert()
		frappe.get_doc({"doctype": "ToDo", "description": "second task", "status": "Cancelled"}).insert()
		source = frappe._dict(
			source_type="DocType",
			reference_doctype="ToDo",
			content_fields="description",
			filters='{"status": "Open"}',
		)
		docs = extract(source)
		self.assertEqual(len(docs), 1)
		self.assertEqual(docs[0].reference_doctype, "ToDo")
		self.assertIn("first task", docs[0].text)
		self.assertTrue(docs[0].reference_name)

	def test_doctype_source_rejects_unknown_field(self):
		source = frappe._dict(
			source_type="DocType",
			reference_doctype="ToDo",
			content_fields="description, not_a_field",
			filters=None,
		)
		with self.assertRaisesRegex(frappe.ValidationError, "Unknown fields"):
			extract(source)

	def test_doctype_source_rejects_invalid_filters(self):
		source = frappe._dict(
			source_type="DocType",
			reference_doctype="ToDo",
			content_fields="description",
			filters="not json",
		)
		with self.assertRaisesRegex(frappe.ValidationError, "valid JSON"):
			extract(source)

	def test_extract_html_strips_tags_and_scripts(self):
		html = "<html><body><script>x=1</script><p>Hi</p> <p>there</p></body></html>"
		self.assertEqual(_extract_html(html), "Hi there")

	def test_extract_xlsx(self):
		import io

		import openpyxl

		workbook = openpyxl.Workbook()
		sheet = workbook.active
		sheet.title = "People"
		sheet.append(["Name", "Age"])
		sheet.append(["Alice", 30])
		buffer = io.BytesIO()
		workbook.save(buffer)
		text = _extract_xlsx(buffer.getvalue())
		self.assertIn("People", text)
		self.assertIn("Alice", text)

	def test_extract_docx(self):
		import io

		from docx import Document

		document = Document()
		document.add_paragraph("First paragraph.")
		document.add_paragraph("")
		document.add_paragraph("Second paragraph.")
		buffer = io.BytesIO()
		document.save(buffer)
		text = _extract_docx(buffer.getvalue())
		self.assertEqual(text, "First paragraph.\n\nSecond paragraph.")

	def test_extract_pdf_blank_page(self):
		import io

		from pypdf import PdfWriter

		writer = PdfWriter()
		writer.add_blank_page(width=72, height=72)
		buffer = io.BytesIO()
		writer.write(buffer)
		self.assertEqual(_extract_pdf(buffer.getvalue()), "")

	def test_validate_public_url_rejects_non_http_scheme(self):
		with self.assertRaisesRegex(frappe.ValidationError, "http"):
			_validate_public_url("ftp://example.com/x")

	def test_validate_public_url_rejects_loopback(self):
		with self.assertRaisesRegex(frappe.ValidationError, "non-public"):
			_validate_public_url("http://127.0.0.1/")

	def test_validate_public_url_rejects_private_ip(self):
		with self.assertRaisesRegex(frappe.ValidationError, "non-public"):
			_validate_public_url("http://10.0.0.5/")

	def test_validate_public_url_rejects_metadata_ip(self):
		with self.assertRaisesRegex(frappe.ValidationError, "non-public"):
			_validate_public_url("http://169.254.169.254/latest/meta-data/")

	def test_validate_public_url_accepts_public_host(self):
		with patch("socket.getaddrinfo", return_value=[(2, 1, 6, "", ("93.184.216.34", 0))]):
			_validate_public_url("https://example.com/page")


class TestIngest(IntegrationTestCase):
	def setUp(self):
		store.drop_table()
		self.model = _make_model()
		_set_settings(
			embedding_model=self.model.name,
			embedding_dimension=DIM,
			chunk_size=60,
			chunk_overlap=10,
		)
		self.kb = frappe.get_doc({"doctype": "AI Knowledge Base", "title": "Ingest KB"}).insert()

	def tearDown(self):
		store.drop_table()
		frappe.db.rollback()

	def _make_source(self, **values):
		values.setdefault("source_type", "Text")
		values.setdefault("title", "Ingest Source")
		with patch("flow.knowledge.ingest.enqueue_ingestion"):
			return frappe.get_doc(
				{"doctype": "AI Knowledge Source", "knowledge_base": self.kb.name, **values}
			).insert()

	def _ingest(self, source_name):
		def fake(input, **kwargs):
			return _embedding_response([[0.1, 0.2, 0.3, 0.4]] * len(input))

		with (
			patch("litellm.embedding", side_effect=fake),
			patch.object(frappe.db, "commit"),
			patch.object(frappe.db, "rollback"),
		):
			ingest_source(source_name)

	def _chunks(self, source_name):
		return frappe.get_all(
			"AI Knowledge Chunk",
			filters={"source": source_name},
			fields=["name", "chunk_index", "knowledge_base"],
			order_by="chunk_index asc",
		)

	def _lance_count(self):
		return store._open_table().count_rows() if store.table_exists() else 0

	def test_text_source_ingests_to_both_stores(self):
		source = self._make_source(content="Wifi keeps dropping on the office network every afternoon. " * 4)
		self._ingest(source.name)

		source.reload()
		self.assertEqual(source.status, "Completed")

		chunks = self._chunks(source.name)
		self.assertGreater(len(chunks), 1)
		self.assertEqual(source.chunk_count, len(chunks))
		self.assertEqual([c["chunk_index"] for c in chunks], list(range(len(chunks))))
		self.assertTrue(all(c["knowledge_base"] == self.kb.name for c in chunks))

		table = store._open_table()
		self.assertEqual(table.count_rows(), len(chunks))
		lance_ids = set(table.to_arrow().column("id").to_pylist())
		self.assertEqual(lance_ids, {int(c["name"]) for c in chunks})

	def test_resync_is_idempotent(self):
		source = self._make_source(content="Reset your password from the account settings page. " * 4)
		self._ingest(source.name)
		first = self._chunks(source.name)
		first_lance = self._lance_count()

		self._ingest(source.name)
		second = self._chunks(source.name)

		self.assertEqual(len(first), len(second))
		self.assertEqual(self._lance_count(), first_lance)
		self.assertEqual(self._lance_count(), len(second))
		self.assertEqual(len({c["name"] for c in second}), len(second))

	def test_trash_cleans_both_stores(self):
		source = self._make_source(content="VPN setup guide for remote employees. " * 4)
		self._ingest(source.name)
		self.assertGreater(self._lance_count(), 0)

		source.delete()

		self.assertEqual(frappe.db.count("AI Knowledge Chunk", {"source": source.name}), 0)
		self.assertEqual(self._lance_count(), 0)

	def test_empty_source_completes_with_zero_chunks(self):
		source = self._make_source(content="   ")
		self._ingest(source.name)

		source.reload()
		self.assertEqual(source.status, "Completed")
		self.assertEqual(source.chunk_count, 0)
		self.assertEqual(self._chunks(source.name), [])

	def test_failed_source_marks_failed_and_raises(self):
		source = self._make_source(source_type="URL", url="http://127.0.0.1/secret")
		with self.assertRaises(frappe.ValidationError):
			self._ingest(source.name)

		source.reload()
		self.assertEqual(source.status, "Failed")
		self.assertIn("non-public", source.error_log or "")

	def test_purge_source_is_safe_when_empty(self):
		source = self._make_source(content="never ingested")
		purge_source(source.name)
		self.assertEqual(frappe.db.count("AI Knowledge Chunk", {"source": source.name}), 0)

	def test_after_insert_enqueues_ingestion(self):
		with patch("flow.knowledge.ingest.enqueue_ingestion") as mock:
			source = frappe.get_doc(
				{
					"doctype": "AI Knowledge Source",
					"knowledge_base": self.kb.name,
					"source_type": "Text",
					"title": "Auto Ingest",
					"content": "indexed on creation",
				}
			).insert()
		mock.assert_called_once_with(source.name)


class TestDoctypeSync(IntegrationTestCase):
	def setUp(self):
		store.drop_table()
		self.model = _make_model()
		_set_settings(
			embedding_model=self.model.name,
			embedding_dimension=DIM,
			chunk_size=200,
			chunk_overlap=20,
		)
		self.kb = frappe.get_doc({"doctype": "AI Knowledge Base", "title": "Sync KB"}).insert()

	def tearDown(self):
		store.drop_table()
		frappe.db.rollback()

	def _todo(self, description, **values):
		return frappe.get_doc({"doctype": "ToDo", "description": description, **values}).insert()

	def _source(self, filters=None):
		filters = filters or {"description": ["like", "SYNCTEST%"]}
		with patch("flow.knowledge.ingest.enqueue_ingestion"):
			return frappe.get_doc(
				{
					"doctype": "AI Knowledge Source",
					"knowledge_base": self.kb.name,
					"source_type": "DocType",
					"title": "ToDo Source",
					"reference_doctype": "ToDo",
					"content_fields": "description",
					"filters": json.dumps(filters),
				}
			).insert()

	def _sync(self, source_name):
		def fake(input, **kwargs):
			return _embedding_response([[0.1, 0.2, 0.3, 0.4]] * len(input))

		embed = patch("litellm.embedding", side_effect=fake)
		with embed as mock, patch.object(frappe.db, "commit"), patch.object(frappe.db, "rollback"):
			ingest_source(source_name)
		return mock

	def _chunks(self, source_name, reference_name=None):
		filters = {"source": source_name}
		if reference_name:
			filters["reference_name"] = reference_name
		return frappe.get_all(
			"AI Knowledge Chunk",
			filters=filters,
			fields=["name", "reference_name", "content_hash"],
			order_by="name asc",
		)

	def _lance_ids(self):
		if not store.table_exists():
			return set()
		return set(store._open_table().to_arrow().column("id").to_pylist())

	def test_backfill_indexes_filtered_rows(self):
		self._todo("SYNCTEST wifi keeps dropping on the office network")
		self._todo("SYNCTEST vpn will not connect after the update")
		self._todo("UNRELATED should not be indexed")
		src = self._source()
		self._sync(src.name)

		src.reload()
		self.assertEqual(src.status, "Completed")
		self.assertIsNotNone(src.last_synced_at)
		chunks = self._chunks(src.name)
		self.assertEqual(len({c["reference_name"] for c in chunks}), 2)
		self.assertTrue(all(c["content_hash"] for c in chunks))
		self.assertEqual(self._lance_ids(), {int(c["name"]) for c in chunks})

	def test_insert_adds_only_new_row(self):
		self._todo("SYNCTEST first ticket about printers")
		src = self._source()
		self._sync(src.name)
		before = {c["name"] for c in self._chunks(src.name)}

		self._todo("SYNCTEST second ticket about scanners")
		self._sync(src.name)

		chunks = self._chunks(src.name)
		self.assertEqual(len({c["reference_name"] for c in chunks}), 2)
		self.assertTrue(before.issubset({c["name"] for c in chunks}))
		self.assertEqual(self._lance_ids(), {int(c["name"]) for c in chunks})

	def test_update_content_replaces_chunks(self):
		todo = self._todo("SYNCTEST wifi keeps dropping")
		src = self._source()
		self._sync(src.name)
		before = self._chunks(src.name, todo.name)
		old_names = {c["name"] for c in before}

		frappe.db.set_value(
			"ToDo",
			todo.name,
			"description",
			"SYNCTEST wifi fixed by replacing the router",
			update_modified=True,
		)
		self._sync(src.name)

		after = self._chunks(src.name, todo.name)
		self.assertNotEqual(after[0]["content_hash"], before[0]["content_hash"])
		self.assertTrue(old_names.isdisjoint({c["name"] for c in after}))
		self.assertEqual(self._lance_ids(), {int(c["name"]) for c in after})

	def test_noncontent_change_is_skipped(self):
		todo = self._todo("SYNCTEST vpn issue", priority="Low")
		src = self._source()
		self._sync(src.name)
		before = {c["name"]: c["content_hash"] for c in self._chunks(src.name, todo.name)}

		frappe.db.set_value("ToDo", todo.name, "priority", "High", update_modified=True)
		mock = self._sync(src.name)

		after = {c["name"]: c["content_hash"] for c in self._chunks(src.name, todo.name)}
		self.assertEqual(after, before)
		self.assertEqual(mock.call_count, 0)

	def test_filter_exit_removes_chunks(self):
		todo = self._todo("SYNCTEST printer offline", status="Open")
		src = self._source(filters={"description": ["like", "SYNCTEST%"], "status": "Open"})
		self._sync(src.name)
		self.assertEqual(len(self._chunks(src.name, todo.name)), 1)

		frappe.db.set_value("ToDo", todo.name, "status", "Cancelled", update_modified=True)
		mock = self._sync(src.name)

		self.assertEqual(self._chunks(src.name, todo.name), [])
		self.assertEqual(self._lance_ids(), set())
		self.assertEqual(mock.call_count, 0)

	def test_delete_removes_chunks(self):
		keep = self._todo("SYNCTEST keep this one")
		drop = self._todo("SYNCTEST delete this one")
		src = self._source()
		self._sync(src.name)
		self.assertEqual(len({c["reference_name"] for c in self._chunks(src.name)}), 2)

		frappe.delete_doc("ToDo", drop.name)
		self._sync(src.name)

		self.assertEqual({c["reference_name"] for c in self._chunks(src.name)}, {keep.name})
		self.assertEqual(self._lance_ids(), {int(c["name"]) for c in self._chunks(src.name)})

	def test_reconcile_purges_chunks_for_tombstoneless_deletes(self):
		keep = self._todo("SYNCTEST keep this one")
		drop = self._todo("SYNCTEST delete this one")
		src = self._source()
		self._sync(src.name)
		self.assertEqual(len({c["reference_name"] for c in self._chunks(src.name)}), 2)

		frappe.db.delete("ToDo", {"name": drop.name})  # raw delete: no Deleted Document, sweep can't see it
		self._sync(src.name)
		self.assertEqual(len({c["reference_name"] for c in self._chunks(src.name)}), 2)  # still orphaned

		reconcile_source(src.name)

		self.assertEqual({c["reference_name"] for c in self._chunks(src.name)}, {keep.name})
		self.assertEqual(self._lance_ids(), {int(c["name"]) for c in self._chunks(src.name)})

	def test_sync_due_sources_enqueues_only_auto_sync_doctypes(self):
		auto = self._source()
		auto.db_set("auto_sync", 1, update_modified=False)
		manual = self._source()
		with patch("flow.knowledge.ingest.enqueue_ingestion"):
			text = frappe.get_doc(
				{
					"doctype": "AI Knowledge Source",
					"knowledge_base": self.kb.name,
					"source_type": "Text",
					"title": "Text Source",
					"content": "not a doctype",
					"auto_sync": 1,
				}
			).insert()

		with patch("flow.knowledge.ingest.enqueue_ingestion") as mock:
			sync_due_sources()
		enqueued = {call.args[0] for call in mock.call_args_list}

		self.assertIn(auto.name, enqueued)
		self.assertNotIn(manual.name, enqueued)
		self.assertNotIn(text.name, enqueued)


class TestRetriever(IntegrationTestCase):
	def setUp(self):
		store.drop_table()
		self.model = _make_model()
		_set_settings(
			embedding_model=self.model.name,
			embedding_dimension=DIM,
			chunk_size=200,
			chunk_overlap=20,
		)
		self.kb = frappe.get_doc({"doctype": "AI Knowledge Base", "title": "Retrieve KB"}).insert()

	def tearDown(self):
		store.drop_table()
		frappe.db.rollback()

	def _fake_embed(self, input, **kwargs):
		return _embedding_response([[0.1, 0.2, 0.3, 0.4]] * len(input))

	def _ingest(self, source_name):
		with (
			patch("litellm.embedding", side_effect=self._fake_embed),
			patch.object(frappe.db, "commit"),
			patch.object(frappe.db, "rollback"),
		):
			ingest_source(source_name)

	def _text_source(self, kb, content, title="Doc"):
		with patch("flow.knowledge.ingest.enqueue_ingestion"):
			source = frappe.get_doc(
				{
					"doctype": "AI Knowledge Source",
					"knowledge_base": kb,
					"source_type": "Text",
					"title": title,
					"content": content,
				}
			).insert()
		self._ingest(source.name)
		return source

	def _retrieve(self, query, **kwargs):
		with patch("litellm.embedding", side_effect=self._fake_embed):
			return retrieve(query, **kwargs)

	def test_retrieve_returns_scoped_chunk_content(self):
		source = self._text_source(self.kb.name, "laptop setup guide for new employees. " * 4)
		results = self._retrieve("laptop", kbs=[self.kb.name])
		self.assertTrue(results)
		self.assertTrue(all("laptop" in r["content"] for r in results))
		self.assertTrue(all(r["source"] == source.name for r in results))

	def test_retrieve_requires_a_knowledge_base(self):
		with self.assertRaises(frappe.ValidationError):
			retrieve("laptop", kbs=[])

	def test_retrieve_blank_query_returns_empty(self):
		self._text_source(self.kb.name, "laptop setup guide. " * 4)
		self.assertEqual(self._retrieve("   ", kbs=[self.kb.name]), [])

	def test_retrieve_excludes_other_knowledge_bases(self):
		other = frappe.get_doc({"doctype": "AI Knowledge Base", "title": "Other KB"}).insert()
		mine = self._text_source(self.kb.name, "laptop onboarding steps. " * 4)
		self._text_source(other.name, "laptop onboarding steps. " * 4, title="Other Doc")
		results = self._retrieve("laptop", kbs=[self.kb.name])
		self.assertTrue(results)
		self.assertTrue(all(r["source"] == mine.name for r in results))


class TestKnowledgeBuilder(IntegrationTestCase):
	def setUp(self):
		store.drop_table()
		self.model = _make_model()
		_set_settings(
			embedding_model=self.model.name,
			embedding_dimension=DIM,
			chunk_size=200,
			chunk_overlap=20,
		)

	def tearDown(self):
		store.drop_table()
		frappe.db.rollback()

	def _fake_embed(self, input, **kwargs):
		return _embedding_response([[0.1, 0.2, 0.3, 0.4]] * len(input))

	def _add_text(self, kb, content, **kwargs):
		with (
			patch("litellm.embedding", side_effect=self._fake_embed),
			patch.object(frappe.db, "commit"),
			patch.object(frappe.db, "rollback"),
		):
			return kb.add_text(content, **kwargs)

	def test_get_or_create_is_idempotent(self):
		first = Knowledge("Shared KB")
		second = Knowledge("Shared KB")
		self.assertEqual(first.name, second.name)
		self.assertEqual(frappe.db.count("AI Knowledge Base", {"title": "Shared KB"}), 1)

	def test_add_text_makes_content_searchable(self):
		kb = Knowledge("Builder KB")
		source_name = self._add_text(kb, "laptop setup guide for new employees. " * 4)

		source = frappe.get_doc("AI Knowledge Source", source_name)
		self.assertEqual(source.source_type, "Text")
		self.assertEqual(source.knowledge_base, kb.name)

		with patch("litellm.embedding", side_effect=self._fake_embed):
			results = retrieve("laptop", kbs=[kb.name])
		self.assertTrue(results)
		self.assertTrue(all(r["source"] == source_name for r in results))

	def test_add_local_file_reads_and_indexes_disk_file(self):
		import os
		import tempfile

		fd, path = tempfile.mkstemp(suffix=".txt")
		try:
			with os.fdopen(fd, "w") as handle:
				handle.write("laptop provisioning runbook for it staff. " * 4)
			kb = Knowledge("Local File KB")
			with (
				patch("litellm.embedding", side_effect=self._fake_embed),
				patch.object(frappe.db, "commit"),
				patch.object(frappe.db, "rollback"),
			):
				source_name = kb.add_local_file(path)
		finally:
			os.unlink(path)

		source = frappe.get_doc("AI Knowledge Source", source_name)
		self.assertEqual(source.source_type, "Text")
		self.assertEqual(source.title, os.path.basename(path))

		with patch("litellm.embedding", side_effect=self._fake_embed):
			results = retrieve("laptop", kbs=[kb.name])
		self.assertTrue(results)
		self.assertTrue(all(r["source"] == source_name for r in results))

	def test_add_text_does_not_enqueue_async_ingestion(self):
		kb = Knowledge("Sync KB")
		with patch("flow.knowledge.ingest.enqueue_ingestion") as enqueue:
			self._add_text(kb, "synchronous ingestion content. " * 4)
		enqueue.assert_not_called()

	def test_add_doctype_translates_fields_and_filters(self):
		kb = Knowledge("DocType KB")
		with patch("flow.knowledge.ingest.ingest_source") as ingest:
			source_name = kb.add_doctype(
				"ToDo",
				content_fields=["description", "status"],
				filters={"status": "Open"},
			)

		source = frappe.get_doc("AI Knowledge Source", source_name)
		self.assertEqual(source.source_type, "DocType")
		self.assertEqual(source.reference_doctype, "ToDo")
		self.assertEqual(source.content_fields, "description, status")
		self.assertEqual(json.loads(source.filters), {"status": "Open"})
		ingest.assert_called_once_with(source_name)


class TestAgentKnowledge(IntegrationTestCase):
	def tearDown(self):
		frappe.db.rollback()

	def _agent(self, knowledge):
		from flow.lib.agent import Agent
		from flow.lib.model import Model

		agent = Agent(model=Model(model_id="openai/gpt-4o-mini"), knowledge=knowledge)
		return next(t for t in agent.tools if t.name == "search_knowledge")

	def test_knowledge_binds_search_tool_exposing_only_query(self):
		search = self._agent(Knowledge("Agent KB"))
		self.assertEqual(set(search.parameters["properties"]), {"query"})

	def test_knowledge_scopes_search_to_given_bases(self):
		kb = Knowledge("Scoped KB")
		search = self._agent([kb])
		with patch("flow.knowledge.retriever.retrieve", return_value=[]) as mock:
			search(query="hi")
		mock.assert_called_once_with("hi", kbs=[kb.name])

	def test_multiple_knowledge_bases_are_all_scoped(self):
		first = Knowledge("KB One")
		second = Knowledge("KB Two")
		search = self._agent([first, second])
		with patch("flow.knowledge.retriever.retrieve", return_value=[]) as mock:
			search(query="hi")
		mock.assert_called_once_with("hi", kbs=[first.name, second.name])
